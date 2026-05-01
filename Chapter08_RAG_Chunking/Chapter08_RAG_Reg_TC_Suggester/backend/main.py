import os
import glob
import json
import requests
import pandas as pd
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import chromadb
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, UnstructuredExcelLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from groq import Groq
from docx import Document
from fpdf import FPDF

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
CHROMA_DB_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "chroma_db")

# Ensure directories exist
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(CHROMA_DB_DIR, exist_ok=True)

chroma_client = chromadb.PersistentClient(path=CHROMA_DB_DIR)

try:
    embeddings = HuggingFaceEmbeddings(
        model_name="nomic-ai/nomic-embed-text-v1.5",
        model_kwargs={'trust_remote_code': True}
    )
except Exception as e:
    print(f"Error loading embeddings: {e}")
    embeddings = None

collection = chroma_client.get_or_create_collection(name="vwo_test_cases")

class GenerateRequest(BaseModel):
    jira_id: str
    jira_domain: str
    jira_email: str
    jira_token: str
    groq_api_key: str
    output_format: str  # "pdf", "docx", "xlsx"

class QueryRequest(BaseModel):
    question: str
    groq_api_key: str

@app.post("/generate_tc")
async def generate_test_cases(req: GenerateRequest):
    # 1. Fetch Jira Issue
    jira_url = f"https://{req.jira_domain}.atlassian.net/rest/api/2/issue/{req.jira_id}"
    auth = (req.jira_email, req.jira_token)
    
    response = requests.get(jira_url, auth=auth, headers={"Accept": "application/json"})
    
    if response.status_code != 200:
        raise HTTPException(status_code=400, detail=f"Failed to fetch Jira Issue. Check credentials. Status: {response.status_code}")
    
    issue_data = response.json()
    summary = issue_data.get("fields", {}).get("summary", "No Summary")
    description = issue_data.get("fields", {}).get("description", "No Description")
    
    # 2. Call Groq API to generate test cases in JSON format
    client = Groq(api_key=req.groq_api_key)
    
    system_prompt = (
        "You are an expert QA Engineer. Generate comprehensive test cases for the provided Jira Issue. "
        "Aim to cover functional, non-functional, security, edge cases, and boundary values. "
        "You MUST return ONLY a valid JSON array of objects. Do not include markdown formatting or extra text. "
        "Format: [{\"tc_id\": \"TC-01\", \"title\": \"Verify X\", \"steps\": \"1. Do A 2. Do B\", \"expected\": \"Result Y\"}]"
    )
    
    user_prompt = f"Jira ID: {req.jira_id}\nSummary: {summary}\nDescription: {description}"
    
    content = ""
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.3,
            max_tokens=2000
        )
        
        content = chat_completion.choices[0].message.content
        # Clean JSON if there are markdown ticks
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].split("```")[0].strip()
            
        tc_list = json.loads(content)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM Generation Error or JSON parse error: {str(e)}\nOutput was: {content}")

    # 3. Save as requested format
    file_name = f"{req.jira_id}_TestCases.{req.output_format}"
    file_path = os.path.join(DATA_DIR, file_name)
    
    if req.output_format == "xlsx":
        df = pd.DataFrame(tc_list)
        df.to_excel(file_path, index=False)
    
    elif req.output_format == "docx":
        doc = Document()
        doc.add_heading(f"Test Cases for {req.jira_id}: {summary}", 0)
        for tc in tc_list:
            doc.add_heading(f"{tc.get('tc_id', '')} - {tc.get('title', '')}", level=1)
            doc.add_paragraph(f"Steps: {tc.get('steps', '')}")
            doc.add_paragraph(f"Expected Result: {tc.get('expected', '')}")
            doc.add_paragraph("-" * 20)
        doc.save(file_path)
        
    elif req.output_format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add a title
        pdf.set_font("Arial", 'B', 16)
        # Using built-in replace to avoid unicode encoding errors in FPDF basic
        safe_summary = summary.encode('ascii', 'ignore').decode('ascii')
        pdf.cell(200, 10, txt=f"Test Cases for {req.jira_id}: {safe_summary}", ln=True, align='C')
        pdf.ln(10)
        
        pdf.set_font("Arial", size=10)
        for tc in tc_list:
            safe_id = str(tc.get('tc_id', '')).encode('ascii', 'ignore').decode('ascii')
            safe_title = str(tc.get('title', '')).encode('ascii', 'ignore').decode('ascii')
            safe_steps = str(tc.get('steps', '')).encode('ascii', 'ignore').decode('ascii')
            safe_expected = str(tc.get('expected', '')).encode('ascii', 'ignore').decode('ascii')
            
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(200, 10, txt=f"{safe_id} - {safe_title}", ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 7, txt=f"Steps: {safe_steps}")
            pdf.multi_cell(0, 7, txt=f"Expected: {safe_expected}")
            pdf.cell(200, 10, txt="-----------------------", ln=True)
            
        pdf.output(file_path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported output format. Choose pdf, docx, or xlsx.")
        
    return {"message": "Success", "file_name": file_name, "total_generated": len(tc_list)}

@app.post("/ingest")
async def ingest_documents():
    """Reads PDFs, DOCX, and XLSX files from the data folder, chunks them, and stores in ChromaDB"""
    if not embeddings:
        raise HTTPException(status_code=500, detail="Embeddings model not loaded properly.")
    
    pdf_files = glob.glob(os.path.join(DATA_DIR, "*.pdf"))
    docx_files = glob.glob(os.path.join(DATA_DIR, "*.docx"))
    xlsx_files = glob.glob(os.path.join(DATA_DIR, "*.xlsx"))
    all_files = pdf_files + docx_files + xlsx_files
    
    if not all_files:
        raise HTTPException(status_code=404, detail="No files found in the 'data' folder.")
    
    all_chunks = []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    
    for file_path in all_files:
        try:
            if file_path.endswith('.pdf'):
                loader = PyPDFLoader(file_path)
            elif file_path.endswith('.docx'):
                loader = Docx2txtLoader(file_path)
            elif file_path.endswith('.xlsx'):
                loader = UnstructuredExcelLoader(file_path)
            else:
                continue
                
            docs = loader.load()
            chunks = text_splitter.split_documents(docs)
            
            for i, chunk in enumerate(chunks):
                chunk.metadata["source"] = os.path.basename(file_path)
                chunk.metadata["chunk_id"] = f"{os.path.basename(file_path)}_chunk_{i}"
                all_chunks.append(chunk)
        except Exception as e:
            print(f"Failed to process {file_path}: {e}")
            
    if not all_chunks:
        return {"message": "No text found in documents to chunk."}
        
    documents = [chunk.page_content for chunk in all_chunks]
    metadatas = [chunk.metadata for chunk in all_chunks]
    ids = [chunk.metadata["chunk_id"] for chunk in all_chunks]
    
    embeds = embeddings.embed_documents(documents)
    
    collection.upsert(
        documents=documents,
        embeddings=embeds,
        metadatas=metadatas,
        ids=ids
    )
    
    return {
        "message": "Ingestion successful!", 
        "total_chunks_stored": len(all_chunks),
        "files_processed": [os.path.basename(f) for f in all_files]
    }

@app.get("/database")
async def get_database():
    results = collection.get()
    return {
        "ids": results.get("ids", []),
        "documents": results.get("documents", []),
        "metadatas": results.get("metadatas", [])
    }

@app.post("/query")
async def query_rag(request: QueryRequest):
    if not request.groq_api_key:
        raise HTTPException(status_code=400, detail="Groq API key is required.")
        
    question_embedding = embeddings.embed_query(request.question)
    
    results = collection.query(
        query_embeddings=[question_embedding],
        n_results=4
    )
    
    if not results["documents"] or not results["documents"][0]:
        return {"answer": "No relevant documents found in the database. Please ingest documents first.", "retrieved_chunks": []}
        
    retrieved_chunks = results["documents"][0]
    retrieved_metadatas = results["metadatas"][0]
    retrieved_ids = results["ids"][0]
    
    context = "\n\n".join([f"Chunk ID: {retrieved_ids[i]}\n{doc}" for i, doc in enumerate(retrieved_chunks)])
    
    client = Groq(api_key=request.groq_api_key)
    
    system_prompt = (
        "You are an expert QA Assistant. Answer the question based ONLY on the provided context. "
        "If the context does not contain the answer, say 'I cannot answer this based on the provided documents.'\n\n"
        f"Context:\n{context}"
    )
    
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": request.question}
            ],
            model="llama-3.1-8b-instant", 
            temperature=0.2,
        )
        answer = chat_completion.choices[0].message.content
        
        return {
            "answer": answer,
            "retrieved_chunks": [
                {"id": retrieved_ids[i], "content": retrieved_chunks[i], "metadata": retrieved_metadatas[i]}
                for i in range(len(retrieved_chunks))
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Groq API Error: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)
